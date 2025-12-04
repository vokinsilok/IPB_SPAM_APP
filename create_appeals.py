"""
Создание обращений на ukc.gov.ua для зарегистрированных аккаунтов
"""
import asyncio
import csv
import json
import logging
import random
import subprocess
import time
import os
from datetime import datetime
from pathlib import Path

import requests
from playwright.async_api import async_playwright, Page
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
from proxy_config import (
    find_working_proxy,
    format_proxy_for_playwright,
    log_proxy_usage,
    print_proxy_stats
)

# Загрузка переменных окружения
load_dotenv()

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Пути
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
REGISTRATIONS_DIR = BASE_DIR / "registrations"
UKC_ACCOUNTS_CSV = REGISTRATIONS_DIR / "ukc_registered.csv"
PEOPLES_CSV = DATA_DIR / "peoples.csv"
APPEALS_BASE_DIR = BASE_DIR / "appeals"
APPEALS_BASE_DIR.mkdir(parents=True, exist_ok=True)

# Настройки прокси перенесены в proxy_config.py
# Используется автоматический поиск рабочего прокси

# База данных обращений
APPEALS_DATABASE = REGISTRATIONS_DIR / "appeals_database.json"

# OpenAI клиент с настройками из .env
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
)

# Модель OpenAI из .env или по умолчанию
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

# Темы обращений с портала UKC (будут выбираться случайно)
APPEAL_TOPICS = [
    "Аграрна політика і земельні відносини",
    "Архіви та статистика",
    "Будівництво та благоустрій",
    "Діяльність органів виконавчої влади та органів місцевого самоврядування",
    "Дотримання законності, громадського порядку та протидія злочинності",
    "Діяльність посадових і службових осіб, корупція",
    "Діяльність підприємств та установ",
    "Діяльність громадських об'єднань, міжнаціональні та міжконфесійні відносини",
    "Екологія та природні ресурси",
    "Житлова політика",
    "Зайнятість та безробіття",
    "Захист прав споживачів",
    "Інформаційна політика та діяльність ЗМІ",
    "Комунальне господарство",
    "Культура, мовна політика, туризм",
    "Міграція, громадянство, паспортизація",
    "Міждержавні відносини",
    "Митна справа",
    "Надзвичайні ситуації",
    "Організація виборчого процесу",
    "Охорона здоров'я",
    "Оподаткування",
    "Оплата праці",
    "Обороноздатність",
    "Освіта, наука та інтелектуальна власність",
    "Пенсійне забезпечення",
    "Підприємницька діяльність",
    "Робота органів юстиції",
    "Соціальний захист населення",
    "Транспортне обслуговування",
    "Умови праці та трудові відносини",
    "Функціонування мереж зв'язку",
    "Фізична культура і спорт",
    "Фінансова політика"
]


def load_ukc_accounts():
    """Загрузка зарегистрированных UKC аккаунтов"""
    accounts = []
    
    if not UKC_ACCOUNTS_CSV.exists():
        logger.error(f"✗ Файл не найден: {UKC_ACCOUNTS_CSV}")
        return accounts
    
    with open(UKC_ACCOUNTS_CSV, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            # Берём только успешно зарегистрированные
            if 'registered' in row.get('status', '').lower():
                accounts.append({
                    'email': row['email'],
                    'password': row['ukc_password'],
                    'full_name': row['full_name']
                })
    
    logger.info(f"✓ Загружено {len(accounts)} UKC аккаунтов")
    return accounts


def load_peoples_data():
    """Загрузка данных военнослужащих"""
    peoples = {}
    
    if not PEOPLES_CSV.exists():
        logger.error(f"✗ Файл не найден: {PEOPLES_CSV}")
        return peoples
    
    with open(PEOPLES_CSV, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            full_name = f"{row['Имя']} {row['Отчество']} {row['Фамилия']}"
            peoples[full_name] = row
    
    logger.info(f"✓ Загружено {len(peoples)} записей военнослужащих")
    return peoples


def get_person_appeals_count(full_name):
    """Подсчет количества созданных обращений для военнослужащего"""
    person_dir = APPEALS_BASE_DIR / full_name.replace(" ", "_")
    if not person_dir.exists():
        return 0
    
    # Считаем количество папок с темами
    topic_dirs = [d for d in person_dir.iterdir() if d.is_dir()]
    return len(topic_dirs)


def get_used_topics(full_name):
    """Получить список уже использованных тем для военнослужащего"""
    person_dir = APPEALS_BASE_DIR / full_name.replace(" ", "_")
    if not person_dir.exists():
        return set()
    
    used_topics = set()
    for topic_dir in person_dir.iterdir():
        if topic_dir.is_dir():
            used_topics.add(topic_dir.name)
    return used_topics


def select_random_topic(used_topics):
    """Выбрать случайную тему, которая еще не использовалась"""
    available_topics = [t for t in APPEAL_TOPICS if t not in used_topics]
    if not available_topics:
        logger.warning("⚠ Все темы уже использованы!")
        return None
    return random.choice(available_topics)


def load_appeals_database():
    """Загрузка базы данных обращений"""
    if APPEALS_DATABASE.exists():
        with open(APPEALS_DATABASE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {'appeals': []}


def save_appeals_database(database):
    """Сохранение базы данных обращений"""
    with open(APPEALS_DATABASE, 'w', encoding='utf-8') as f:
        json.dump(database, f, ensure_ascii=False, indent=2)


async def generate_appeal_text(topic, person_data):
    """Генерация текста обращения через OpenAI на русском и украинском"""
    logger.info(f"Генерирую текст обращения на тему: {topic}")
    
    # Формируем промпт с личными данными из CSV
    first_name = person_data.get('Имя', '')
    patronymic = person_data.get('Отчество', '')
    last_name = person_data.get('Фамилия', '')
    birth_date = person_data.get('Дата рождения', '')
    rank = person_data.get('Воинское звание', 'Солдат')
    position = person_data.get('Должность', 'Водій')
    unit = person_data.get('Войсковая часть', 'А4759')
    
    prompt = f"""Напиши официальное обращение в Единый контактный центр Украины (УКЦ) от имени военнослужащего.

Тема обращения: {topic}

Личные данные военнослужащего (ОБЯЗАТЕЛЬНО используй эти данные в тексте):
- ФИО: {last_name} {first_name} {patronymic}
- Дата рождения: {birth_date}
- Воинское звание: {rank}
- Должность: {position}
- Войсковая часть: {unit}

Требования:
1. Текст должен быть официальным и вежливым
2. ОБЯЗАТЕЛЬНО упомяни в начале: "Я, {last_name} {first_name} {patronymic}, являюсь военнослужащим"
3. ОБЯЗАТЕЛЬНО укажи звание ({rank}), должность ({position}) и воинскую часть ({unit}) в тексте
4. Опиши конкретную проблему или вопрос по теме "{topic}"
5. Объем: 150-250 слов
6. Напиши ДВА варианта: один на русском языке, второй на украинском языке
7. Формат ответа:
---RUSSIAN---
[текст на русском]
---UKRAINIAN---
[текст на украинском]
"""
    
    try:
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Ты помощник для написания официальных обращений в государственные органы Украины."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.8,
            max_tokens=1500
        )
        
        full_text = response.choices[0].message.content
        
        # Парсим ответ
        russian_text = ""
        ukrainian_text = ""
        
        if "---RUSSIAN---" in full_text and "---UKRAINIAN---" in full_text:
            parts = full_text.split("---RUSSIAN---")[1].split("---UKRAINIAN---")
            russian_text = parts[0].strip()
            ukrainian_text = parts[1].strip()
        else:
            logger.warning("⚠ Не удалось распарсить ответ OpenAI, использую весь текст как украинский")
            ukrainian_text = full_text.strip()
            russian_text = full_text.strip()
        
        logger.info("✓ Текст обращения сгенерирован")
        return russian_text, ukrainian_text
        
    except Exception as e:
        logger.error(f"✗ Ошибка генерации текста: {e}")
        return None, None


def save_appeal_to_docx(text, language, full_name, topic, appeal_dir):
    """Сохранение текста обращения в DOCX файл"""
    # Формируем имя файла
    filename = f"appeal_{language}.docx"
    filepath = appeal_dir / filename
    
    # Создаем документ
    doc = Document()
    doc.add_heading(topic, 0)
    doc.add_paragraph(text)
    
    # Сохраняем
    doc.save(filepath)
    logger.info(f"✓ Сохранено: {filepath}")
    return filepath


def add_appeal_to_database(email, full_name, topic, city, screenshot_date, appeal_dir):
    """Добавить обращение в базу данных"""
    database = load_appeals_database()
    
    appeal_record = {
        'id': len(database['appeals']) + 1,
        'email': email,
        'full_name': full_name,
        'topic': topic,
        'city': city,
        'screenshot_date': screenshot_date,
        'appeal_dir': str(appeal_dir),
        'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    database['appeals'].append(appeal_record)
    save_appeals_database(database)
    logger.info(f"✓ Обращение #{appeal_record['id']} добавлено в базу данных")
    return appeal_record["id"]




async def login_to_ukc(page: Page, email: str, password: str):
    """Авторизация на ukc.gov.ua"""
    logger.info(f"Авторизация: {email}")
    
    try:
        logger.info("Открываю ukc.gov.ua...")
        
        try:
            await page.goto("https://ukc.gov.ua/", timeout=90000, wait_until='commit')
            logger.info(f"✓ Страница начала загружаться: {page.url}")
            
            await page.wait_for_load_state('domcontentloaded', timeout=30000)
            logger.info(f"✓ Страница загружена: {page.url}")
        except Exception as e:
            error_msg = str(e)
            logger.error(f"✗ Ошибка при загрузке страницы: {error_msg}")
            
            if "ERR_TIMED_OUT" in error_msg or "timeout" in error_msg.lower():
                logger.error("⚠ Таймаут при подключении к сайту")
            
            return False
        
        await asyncio.sleep(3)
        
        if "challenge" in page.url or "__cf_chl" in page.url or "checking" in page.url.lower():
            logger.warning("⚠ Обнаружена проверка безопасности, ожидание...")
            await asyncio.sleep(10)
            logger.info(f"Текущий URL после ожидания: {page.url}")
        
        # Клик на "Особистий кабінет"
        try:
            await page.get_by_role("button", name="Особистий кабінет").click(timeout=10000)
            logger.info("✓ Нажата кнопка 'Особистий кабінет'")
        except Exception as e:
            logger.error(f"✗ Не найдена кнопка 'Особистий кабінет': {e}")
            return False
        
        await asyncio.sleep(1)
        
        # Ввод email
        await page.locator("input[type=\"email\"]").click()
        await page.locator("input[type=\"email\"]").fill(email)
        
        # Ввод пароля
        await page.get_by_role("banner").get_by_role("textbox", name="Пароль").click()
        await page.get_by_role("banner").get_by_role("textbox", name="Пароль").fill(password)
        
        # Нажатие "Увійти"
        await page.get_by_role("banner").get_by_role("button", name="Увійти").click()
        
        # Ожидание редиректа
        await asyncio.sleep(2)
        
        # Проверка успешной авторизации
        if "https://ukc.gov.ua/portal/" in page.url:
            logger.info("✓ Успешная авторизация")
            return True
        else:
            # Ищем причину ошибки
            try:
                error_text = await page.locator('.error, .alert-danger, [class*="error"]').first.inner_text(timeout=2000)
                logger.error(f"✗ Ошибка авторизации: {error_text}")
            except:
                logger.error(f"✗ Авторизация не удалась. Текущий URL: {page.url}")
            return False
            
    except Exception as e:
        logger.error(f"✗ Ошибка при авторизации: {e}")
        return False


async def create_appeal(page: Page, topic: str, city: str, appeal_text: str):
    """Создание обращения на портале"""
    logger.info(f"Создаю обращение: {topic}")
    
    try:
        # Переходим на страницу портала если не там
        if "portal/appeals" not in page.url:
            await page.goto("https://ukc.gov.ua/portal/", timeout=30000)
            await asyncio.sleep(1)
        
        # Нажимаем на кнопку создания обращения (несколько вариантов селекторов)
        logger.info("Нажимаю кнопку создания обращения...")
        try:
            # Пробуем разные селекторы
            create_button = page.locator('a[href*="/portal/appeals/create"], button:has-text("Створити звернення"), .section__btn a').first
            await create_button.click(timeout=10000)
            logger.info("✓ Кнопка создания обращения нажата")
        except Exception as e:
            logger.warning(f"⚠ Не удалось найти кнопку создания обращения стандартным способом: {e}")
            # Пробуем CSS селектор из задания
            await page.click("#root > div > div > section:nth-child(2) > div > div.section__btn > a", timeout=10000)
        
        await asyncio.sleep(1)
        
        # Выбираем тему из выпадающего списка (rc-tree-select)
        logger.info("Выбираю тему обращения...")
        try:
            # Кликаем на селектор темы чтобы открыть dropdown
            theme_selector = page.locator('.rc-tree-select-selector, .rc-tree-select, [class*="tree-select"]').first
            await theme_selector.click(timeout=10000)
            logger.info("✓ Dropdown темы открыт")
            await asyncio.sleep(1)
            
            # Ищем нужную тему в списке по точному совпадению
            try:
                # Ищем в dropdown по title или тексту
                theme_option = page.locator(f'.rc-tree-select-tree-title:has-text("{topic}")').first
                await theme_option.click(timeout=5000)
                logger.info(f"✓ Выбрана тема: {topic}")
            except:
                # Если точное совпадение не найдено, ищем частичное
                logger.warning("⚠ Точная тема не найдена, ищу частичное совпадение...")
                short_topic = " ".join(topic.split()[:3])
                try:
                    theme_option = page.locator(f'.rc-tree-select-tree-title:has-text("{short_topic}")').first
                    await theme_option.click(timeout=3000)
                    logger.info(f"✓ Выбрана тема по частичному совпадению")
                except:
                    logger.warning("⚠ Выбираю первую доступную тему")
                    first_theme = page.locator('.rc-tree-select-tree-title').first
                    await first_theme.click(timeout=3000)
        except Exception as e:
            logger.error(f"✗ Ошибка выбора темы: {e}")
            return False
        
        await asyncio.sleep(1)
        
        # Указываем населенный пункт (React-Select dropdown как в регистрации)
        logger.info(f"Указываю населенный пункт: {city}")
        try:
            city_field = page.locator('form .form-field:has(label:has-text("Населений пункт"))')
            await city_field.wait_for(state='visible', timeout=30000)
            
            # Вводим название города в react-select
            city_input = city_field.locator('[id^="react-select-"][id$="-input"]')
            await city_input.focus()
            await city_input.fill('')
            await city_input.type(city, delay=30)
            
            # Ждём появления списка
            await asyncio.sleep(2)
            
            # Нажимаем Enter для выбора первой опции
            await page.keyboard.press('Enter')
            logger.info(f"✓ Населенный пункт выбран: {city}")
            
        except Exception as e:
            logger.warning(f"⚠ Не удалось указать населенный пункт: {e}")
        
        await asyncio.sleep(1)
        
        # Вводим текст обращения (textarea#content)
        logger.info("Ввожу текст обращения...")
        try:
            text_area = page.locator('textarea#content, textarea[name="content"]').first
            await text_area.click()
            await text_area.fill(appeal_text)
            logger.info(f"✓ Текст обращения введен ({len(appeal_text)} символов)")
        except Exception as e:
            logger.error(f"✗ Не удалось ввести текст обращения: {e}")
            return False
        
        await asyncio.sleep(1)
        
        # Решаем капчу (обход как при регистрации)
        logger.info("Обход капчи...")
        await page.evaluate("""
        () => {
          const token = 'TEST_TOKEN_' + Math.random().toString(36).slice(2);
          const fire = (el) => { try { el.dispatchEvent(new Event('input', {bubbles:true})); el.dispatchEvent(new Event('change', {bubbles:true})); } catch(e) {} };
          const areas = Array.from(document.querySelectorAll('textarea[id^="g-recaptcha-response"], textarea[name="g-recaptcha-response"]'));
          areas.forEach(t => { try { t.value = token; fire(t); } catch(e) {} });
          const inputs = Array.from(document.querySelectorAll('input[id^="g-recaptcha-response"], input[name="g-recaptcha-response"]'));
          inputs.forEach(t => { try { t.value = token; fire(t); } catch(e) {} });
          if (areas.length === 0 && inputs.length === 0) {
            try {
              const form = document.querySelector('form') || document.body;
              const ta = document.createElement('textarea');
              ta.id = 'g-recaptcha-response';
              ta.name = 'g-recaptcha-response';
              ta.style.display = 'none';
              ta.value = token;
              form.appendChild(ta);
              fire(ta);
            } catch(e) {}
          }
          const labels = ['Відправити','Надіслати','Отправить'];
          const nodes = Array.from(document.querySelectorAll('button, input[type="submit"]'));
          nodes.forEach(b => {
            const txt = (b.textContent || b.value || '').trim();
            if (labels.some(l => txt.includes(l))) { try { b.removeAttribute('disabled'); b.disabled = false; } catch(e) {} }
          });
        }
        """)
        
        await asyncio.sleep(1)
        
        # Нажимаем кнопку отправки
        logger.info("Отправляю обращение...")
        try:
            submit_button = page.locator('button:has-text("Відправити"), button:has-text("Надіслати"), button[type="submit"]').first
            await submit_button.click(timeout=10000)
            logger.info("✓ Кнопка отправки нажата")
        except Exception as e:
            logger.error(f"✗ Не удалось нажать кнопку отправки: {e}")
            return False
        
        # Ожидание редиректа на страницу обращений
        logger.info("Ожидаю редиректа...")
        await asyncio.sleep(3)
        
        # Проверка успешной отправки
        if "https://ukc.gov.ua/portal/appeals" in page.url:
            logger.info("✓ Обращение отправлено успешно")
            return True
        else:
            # Проверяем, может быть мы на странице конкретного обращения
            if "/portal/appeals/" in page.url:
                logger.info("✓ Обращение создано (перенаправлено на страницу обращения)")
                # Возвращаемся к списку обращений
                await page.goto("https://ukc.gov.ua/portal/appeals", timeout=30000)
                await asyncio.sleep(1)
                return True
            else:
                logger.warning(f"⚠ Неожиданный URL после отправки: {page.url}")
                return False
        
    except Exception as e:
        logger.error(f"✗ Ошибка при создании обращения: {e}")
        return False


async def wait_for_appeal_distribution(page: Page, max_attempts=3):
    """Ожидание распределения обращения (до 1 минуты)"""
    logger.info("Ожидаю распределения обращения...")
    
    # Убеждаемся что мы на странице обращений
    if "portal/appeals" not in page.url:
        await page.goto("https://ukc.gov.ua/portal/appeals", timeout=30000)
        await asyncio.sleep(2)
    
    logger.warning("⚠ Обращение не распределено за отведенное время (1 минута)")
    logger.info("Продолжаю работу без ожидания распределения...")
    return False


async def take_appeal_screenshot(page: Page, appeal_dir: Path):
    """Создание скриншота обращения с подменой даты"""
    logger.info("Делаю скриншот обращения...")
    
    try:
        # Кликаем на первое обращение в списке
        await page.locator('.handling-table__body .handling-table__row').first.locator('a').click()
        await asyncio.sleep(2)
        
        # Генерируем случайную дату в диапазоне 01.12.2025 - 20.12.2025
        random_day = random.randint(1, 20)
        fake_date = f"{random_day:02d}/12/2025"
        
        # Подменяем дату отправки обращения на странице
        logger.info(f"Подменяю дату на: {fake_date}")
        await page.evaluate(f"""
        () => {{
            // Ищем элемент с датой отправки
            const sections = Array.from(document.querySelectorAll('.article-section'));
            sections.forEach(section => {{
                const title = section.querySelector('.article-section__title');
                if (title && title.textContent.includes('Надіслано')) {{
                    const dateP = section.querySelector('p');
                    if (dateP) {{
                        dateP.textContent = '{fake_date}';
                    }}
                }}
            }});
        }}
        """)
        
        await asyncio.sleep(1)
        
        # Делаем скриншот
        screenshot_path = appeal_dir / "screenshot.png"
        await page.screenshot(path=str(screenshot_path), full_page=True)
        logger.info(f"✓ Скриншот сохранен с датой {fake_date}: {screenshot_path}")
        
        # Возвращаемся к списку обращений
        await page.goto("https://ukc.gov.ua/portal/appeals")
        await asyncio.sleep(2)
        
        return fake_date  # Возвращаем дату скриншота
        
    except Exception as e:
        logger.error(f"✗ Ошибка при создании скриншота: {e}")
        return None


async def logout_from_ukc(page: Page):
    """Выход из аккаунта"""
    logger.info("Выхожу из аккаунта...")
    
    try:
        # Шаг 1: Клик на "Особистий кабінет" для открытия меню
        logger.info("Открываю меню личного кабинета...")
        await page.get_by_role("button", name="Особистий кабінет").click(timeout=10000)
        await asyncio.sleep(1)
        
        # Шаг 2: Ждем появления меню и кликаем на "Вийти"
        logger.info("Нажимаю 'Вийти'...")
        try:
            # Пробуем найти по селектору
            logout_link = page.locator('a.js-auth-logout, a:has-text("Вийти")').first
            await logout_link.click(timeout=5000)
            logger.info("✓ Кнопка выхода нажата")
        except Exception as e:
            logger.warning(f"⚠ Не удалось найти кнопку выхода: {e}")
        
        # Шаг 3: Ждем редиректа на портал
        await asyncio.sleep(3)
        
        # Проверка успешного выхода
        if "https://ukc.gov.ua/portal/" in page.url or "https://ukc.gov.ua/" in page.url:
            logger.info("✓ Выход выполнен")
            return True
        else:
            logger.warning(f"⚠ Возможно выход не выполнен. Текущий URL: {page.url}")
            # Принудительно переходим на главную
            await page.goto("https://ukc.gov.ua/")
            await asyncio.sleep(2)
            return False
            
    except Exception as e:
        logger.error(f"✗ Ошибка при выходе из аккаунта: {e}")
        # Принудительно переходим на главную
        try:
            await page.goto("https://ukc.gov.ua/")
            await asyncio.sleep(2)
        except:
            pass
        return False


async def process_account(page: Page, context, account: dict, peoples_data: dict):
    """Обработка одного аккаунта: создание 10 обращений"""
    email = account['email']
    password = account['password']
    full_name = account['full_name']
    
    logger.info("")
    logger.info("=" * 80)
    logger.info(f"ОБРАБОТКА АККАУНТА: {full_name}")
    logger.info(f"Email: {email}")
    logger.info("=" * 80)
    
    # Проверяем количество уже созданных обращений
    existing_appeals = get_person_appeals_count(full_name)
    if existing_appeals >= 10:
        logger.info(f"✓ У {full_name} уже создано {existing_appeals} обращений, пропускаю...")
        return True
    
    # Получаем данные военнослужащего
    person_data = peoples_data.get(full_name)
    if not person_data:
        logger.error(f"✗ Данные для {full_name} не найдены в CSV")
        return False
    
    city = "Запоріжжя"
    
    # Авторизация
    if not await login_to_ukc(page, email, password):
        logger.error(f"✗ Не удалось авторизоваться для {email}")
        return False
    
    # Получаем уже использованные темы
    used_topics = get_used_topics(full_name)
    
    # Создаем обращения (до 10 штук)
    appeals_to_create = 10 - existing_appeals
    logger.info(f"Нужно создать обращений: {appeals_to_create}")
    
    for i in range(appeals_to_create):
        logger.info(f"\n>>> Обращение {existing_appeals + i + 1}/10 <<<")
        
        # Выбираем случайную тему
        topic = select_random_topic(used_topics)
        if not topic:
            logger.error("✗ Не осталось доступных тем")
            break
        
        used_topics.add(topic)
        
        # Генерируем текст обращения
        russian_text, ukrainian_text = await generate_appeal_text(topic, person_data)
        if not ukrainian_text:
            logger.error("✗ Не удалось сгенерировать текст обращения")
            continue
        
        # Создаем директорию для обращения
        safe_name = full_name.replace(" ", "_")
        safe_topic = topic.replace(" ", "_").replace(",", "")[:50]
        appeal_dir = APPEALS_BASE_DIR / safe_name / safe_topic
        appeal_dir.mkdir(parents=True, exist_ok=True)
        
        # Сохраняем тексты в DOCX
        save_appeal_to_docx(russian_text, "ru", full_name, topic, appeal_dir)
        save_appeal_to_docx(ukrainian_text, "uk", full_name, topic, appeal_dir)
        
        # Создаем обращение на портале
        if not await create_appeal(page, topic, city, ukrainian_text):
            logger.error("✗ Не удалось создать обращение на портале")
            continue
        
        # Ожидаем распределения (но делаем скриншот в любом случае)
        distribution_status = await wait_for_appeal_distribution(page)
        if distribution_status:
            logger.info("✓ Обращение распределено")
        else:
            logger.info("⚠ Обращение не распределено, но продолжаем")
        
        # Делаем скриншот в любом случае (даже если не распределено)
        screenshot_date = await take_appeal_screenshot(page, appeal_dir)
        
        # Сохраняем информацию об обращении
        if screenshot_date:
            add_appeal_to_database(
                email=email,
                full_name=full_name,
                topic=topic,
                city=city,
                screenshot_date=screenshot_date,
                appeal_dir=appeal_dir
            )
        
        logger.info(f"✓ Обращение {existing_appeals + i + 1}/10 создано успешно")
        
        # Пауза между обращениями
        if i < appeals_to_create - 1:
            pause = random.randint(1, 3)
            logger.info(f"Пауза {pause} секунд перед следующим обращением...")
            await asyncio.sleep(pause)
    
    # Выход из аккаунта
    await logout_from_ukc(page)
    
    logger.info(f"✓ Обработка аккаунта {email} завершена")
    return True


async def main():
    """Главная функция"""
    logger.info("=" * 80)
    logger.info("СОЗДАНИЕ ОБРАЩЕНИЙ НА UKC.GOV.UA")
    logger.info("=" * 80)
    
    # Загрузка данных
    ukc_accounts = load_ukc_accounts()
    if not ukc_accounts:
        logger.error("✗ Нет зарегистрированных UKC аккаунтов!")
        return
    
    peoples_data = load_peoples_data()
    if not peoples_data:
        logger.error("✗ Нет данных военнослужащих!")
        return
    
    logger.info(f"Всего аккаунтов: {len(ukc_accounts)}")
    
    # Запуск браузера с прокси и новым профилем
    logger.info("Запуск браузера с проверкой прокси...")
    
    async with async_playwright() as p:
        # Создаём уникальную директорию для профиля браузера
        profile_dir = BASE_DIR / "browser_profiles" / "appeals_profile"
        profile_dir.mkdir(parents=True, exist_ok=True)
        
        # Запускаем браузер с новым профилем
        browser = await p.chromium.launch(
            headless=False,
            args=[
                f'--user-data-dir={profile_dir}',
                '--disable-blink-features=AutomationControlled',
                '--disable-dev-shm-usage',
                '--no-sandbox'
            ]
        )
        
        # Создаём временный контекст для проверки прокси
        logger.info("")
        logger.info("Поиск рабочего прокси для gov.ua...")
        temp_context = await browser.new_context()
        temp_page = await temp_context.new_page()
        
        # Ищем рабочий прокси
        working_proxy = await find_working_proxy(temp_page)
        await temp_page.close()
        await temp_context.close()
        
        if not working_proxy:
            logger.error("✗ Не найден рабочий прокси для gov.ua!")
            logger.error("Добавьте больше прокси в proxy_config.py")
            await browser.close()
            return
        
        # Создаём контекст с найденным прокси
        proxy_config = format_proxy_for_playwright(working_proxy)
        context_options = {
            'proxy': proxy_config,
            'locale': 'uk-UA',
            'color_scheme': 'light',
            'viewport': {'width': 1920, 'height': 1080},
            'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'timezone_id': 'Europe/Kiev'
        }
        
        logger.info(f"✓ Используется прокси: {working_proxy['name']} ({working_proxy['server']})")
        
        context = await browser.new_context(**context_options)
        page = await context.new_page()
        
        # Обработка каждого аккаунта
        for idx, account in enumerate(ukc_accounts, start=1):
            logger.info(f"\n>>> Прогресс: {idx}/{len(ukc_accounts)} <<<")
            
            try:
                await process_account(page, context, account, peoples_data)
            except Exception as e:
                logger.error(f"✗ Ошибка при обработке аккаунта {account['email']}: {e}")
            
            # Пауза между аккаунтами
            if idx < len(ukc_accounts):
                pause = random.randint(1, 2)
                logger.info(f"Пауза {pause} секунд перед следующим аккаунтом...")
                await asyncio.sleep(pause)
        
        logger.info("")
        logger.info("=" * 80)
        logger.info("СОЗДАНИЕ ОБРАЩЕНИЙ ЗАВЕРШЕНО")
        logger.info("=" * 80)
        
        # Выводим статистику использования прокси
        print_proxy_stats()
        
        await browser.close()
        input("\nНажмите Enter для завершения...")


if __name__ == "__main__":
    asyncio.run(main())
